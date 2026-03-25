from __future__ import annotations

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_URL_RE = re.compile(r"(?:https?://|www\.)\S+", re.IGNORECASE)
_PHONE_RE = re.compile(r"\+?\d[\d().\-\s]{6,}\d")


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be parsed into usable text."""


def extract_text(raw: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            f"Unsupported file type: {ext or 'unknown'}. Allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if not raw:
        raise DocumentParseError(f"{filename} is empty.")

    try:
        if ext == ".pdf":
            text = _extract_pdf_text(raw, filename)
        elif ext == ".docx":
            text = _extract_docx_text(raw, filename)
        else:
            text = raw.decode("utf-8", errors="ignore")

        text = cleanup_document_text(text)
        if not text:
            raise DocumentParseError(
                f"{filename} produced no readable text. The file may be image-based, encrypted, or malformed."
            )
        return text
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse {filename}: {exc}") from exc


def _extract_pdf_text(raw: bytes, filename: str) -> str:
    reader = _build_pdf_reader(raw, filename)
    pages: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)
        except Exception as exc:
            logger.warning("Skipping unreadable page %s in %s: %s", index, filename, exc)

    text = "\n".join(pages).strip()
    if not text:
        raise DocumentParseError(
            f"{filename} produced no readable PDF text. Try re-exporting the PDF or converting it to DOCX/TXT."
        )
    return text


def _build_pdf_reader(raw: bytes, filename: str):
    last_error: Exception | None = None
    for import_path in ("pypdf", "PyPDF2"):
        try:
            if import_path == "pypdf":
                from pypdf import PdfReader  # type: ignore
            else:
                from PyPDF2 import PdfReader  # type: ignore
            return PdfReader(io.BytesIO(raw))
        except Exception as exc:
            last_error = exc
    raise DocumentParseError(f"Could not open PDF {filename}: {last_error}")


def _extract_docx_text(raw: bytes, filename: str) -> str:
    try:
        import docx
    except Exception as exc:
        raise DocumentParseError("python-docx is not installed.") from exc

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise DocumentParseError(f"Could not open DOCX {filename}: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                tables.append(row_text)

    text = "\n".join(paragraphs + tables).strip()
    if not text:
        raise DocumentParseError(f"{filename} contains no readable DOCX text.")
    return text


def _is_contact_noise_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    marker_count = sum(
        1
        for pattern in (_EMAIL_RE, _URL_RE, _PHONE_RE)
        if pattern.search(stripped)
    )
    lower = stripped.lower()
    if any(token in lower for token in ("linkedin", "github", "portfolio")):
        marker_count += 1
    word_count = len(re.findall(r"\b\w+\b", stripped))
    if marker_count >= 2 and word_count <= 12:
        return True
    if word_count <= 5 and ("|" in stripped or marker_count >= 1):
        return True
    return False


def cleanup_document_text(text: str) -> str:
    text = text.replace("\x00", " ")
    lines = [line.strip() for line in text.splitlines()]
    cleaned: list[str] = []
    seen: set[str] = set()

    for line in lines:
        if not line:
            continue
        if _is_contact_noise_line(line):
            continue
        line = re.sub(r"\s+", " ", line).strip()
        key = line.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(line)

    joined = "\n".join(cleaned)
    joined = re.sub(r"\n{3,}", "\n\n", joined)
    joined = re.sub(r"[ \t]+", " ", joined)
    return joined.strip()


def normalise_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def chunk_text(text: str, chunk_size: int = 180, overlap: int = 30) -> list[str]:
    if not text.strip():
        return []

    blocks = [normalise_whitespace(block) for block in re.split(r"\n\s*\n", text) if block.strip()]
    chunks: list[str] = []

    for block in blocks:
        words = block.split()
        if not words:
            continue
        if len(words) <= chunk_size:
            chunks.append(block)
            continue

        start = 0
        step = max(1, chunk_size - overlap)
        while start < len(words):
            chunk = " ".join(words[start : start + chunk_size]).strip()
            if chunk:
                chunks.append(chunk)
            start += step

    return chunks
